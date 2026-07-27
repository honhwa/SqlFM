# -*- coding: utf-8 -*-
"""将 docs/dev 下的三份 Markdown 开发文档转换为精美 .docx。
支持子集：# 标题(封面)、## H1、### H2、#### H3、表格、```代码块```、
> 引用、- 列表、**加粗**、--- 分页。输出带封面、目录域、页眉页脚页码、表格斑马纹。
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = BASE
FILES = [
    ("需求规格说明书.md", "需求规格说明书.docx"),
    ("软件设计说明书.md", "软件设计说明书.docx"),
    ("测试方案与测试报告.md", "测试方案与测试报告.docx"),
]

CJK = "宋体"
CJK_HEAD = "黑体"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # 深蓝
HEADER_FILL = "1F4E79"
ZEBRA_A = "FFFFFF"
ZEBRA_B = "EEF3F8"


def set_cjk(run, font=CJK, size=None, bold=None, color=None):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcpr.append(shd)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = field_code
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(it); run._r.append(sep); run._r.append(end)
    return run


def add_page_number_footer(section, text):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text + "   第 ")
    set_cjk(r, size=9, color=RGBColor(0x66, 0x66, 0x66))
    add_field(p, "PAGE")
    r2 = p.add_run(" 页 / 共 ")
    set_cjk(r2, size=9, color=RGBColor(0x66, 0x66, 0x66))
    add_field(p, "NUMPAGES")
    r3 = p.add_run(" 页")
    set_cjk(r3, size=9, color=RGBColor(0x66, 0x66, 0x66))


def add_header(section, text):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_cjk(r, size=9, color=RGBColor(0x66, 0x66, 0x66))


def parse_inline(paragraph, text):
    """处理 **加粗** 等行内格式。"""
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            set_cjk(r, bold=True)
        else:
            r = paragraph.add_run(part)
            set_cjk(r)


def build_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    # 默认字体
    normal = doc.styles['Normal']
    normal.font.name = CJK
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)

    # 页边距
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        add_header(s, "SqlFM · 开发文档")
        add_page_number_footer(s, "SqlFM 开发文档")

    cover_title = ""
    cover_sub = ""
    in_code = False
    code_buf = []
    i = 0
    first_h1_done = False

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_cjk(run, font="Consolas", size=9, color=RGBColor(0x33, 0x33, 0x33))
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F4F4F4')
            pPr.append(shd)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(6)
            code_buf = []

    while i < len(lines):
        line = lines[i]
        # 代码块
        if line.strip().startswith('```'):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 封面标题（首个 #）
        if line.startswith('# ') and not cover_title:
            cover_title = line[2:].strip()
            i += 1
            continue
        # 封面副标题（首个 > 引用块，紧跟标题后）
        if line.startswith('> ') and cover_title and not cover_sub and not first_h1_done:
            cover_sub = line[2:].strip()
            i += 1
            continue

        # 分页
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # 标题
        if line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=3)
            for r in p.runs: set_cjk(r, font=CJK_HEAD, size=12, bold=True, color=ACCENT)
            i += 1; continue
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
            for r in p.runs: set_cjk(r, font=CJK_HEAD, size=13, bold=True, color=ACCENT)
            i += 1; continue
        if line.startswith('## '):
            # 首个 H1 前插入目录
            if not first_h1_done:
                # 封面
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tr = cp.add_run(cover_title)
                set_cjk(tr, font=CJK_HEAD, size=26, bold=True, color=ACCENT)
                cp.paragraph_format.space_before = Pt(120)
                sp = doc.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sr = sp.add_run(cover_sub)
                set_cjk(sr, size=12, color=RGBColor(0x55, 0x55, 0x55))
                doc.add_page_break()
                # 目录
                tp = doc.add_paragraph()
                thr = tp.add_run("目录")
                set_cjk(thr, font=CJK_HEAD, size=16, bold=True, color=ACCENT)
                toc_p = doc.add_paragraph()
                add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
                doc.add_page_break()
                first_h1_done = True
            p = doc.add_heading(line[3:].strip(), level=1)
            for r in p.runs: set_cjk(r, font=CJK_HEAD, size=15, bold=True, color=ACCENT)
            i += 1; continue

        # 引用
        if line.startswith('> '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            set_cjk(r, size=9, color=RGBColor(0x66, 0x66, 0x66))
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1; continue

        # 表格
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|') and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
            # 收集表头 + 数据行
            header_cells = [c.strip() for c in line.strip().strip('|').split('|')]
            data_rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith('|'):
                data_rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                j += 1
            tbl = doc.add_table(rows=1, cols=len(header_cells))
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            for k, h in enumerate(header_cells):
                hdr[k].text = ''
                pr = hdr[k].paragraphs[0].add_run(h)
                set_cjk(pr, bold=True, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF))
                shade_cell(hdr[k], HEADER_FILL)
            for ri, row in enumerate(data_rows):
                cells = tbl.add_row().cells
                fill = ZEBRA_A if ri % 2 == 0 else ZEBRA_B
                for k, val in enumerate(row):
                    cells[k].text = ''
                    parse_inline(cells[k].paragraphs[0], val)
                    for rr in cells[k].paragraphs[0].runs:
                        set_cjk(rr, size=9.5)
                    shade_cell(cells[k], fill)
            # 列宽自适应
            for row in tbl.rows:
                for c in row.cells:
                    c.width = Inches(6.5 / max(len(header_cells), 1))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i = j
            continue

        # 列表
        if re.match(r'^\s*[-*] ', line):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, re.sub(r'^\s*[-*] ', '', line))
            for r in p.runs: set_cjk(r)
            i += 1; continue

        # 空行
        if not line.strip():
            i += 1; continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline(p, line.strip())
        for r in p.runs: set_cjk(r)
        p.paragraph_format.space_after = Pt(4)
        i += 1

    flush_code()
    out_path = os.path.join(OUT, docx_path)
    doc.save(out_path)
    print("生成:", out_path)


if __name__ == '__main__':
    for md, docx in FILES:
        build_docx(os.path.join(BASE, md), docx)
    print("全部完成。")
